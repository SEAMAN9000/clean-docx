# -*- coding: utf-8 -*-
"""中文文档排版工具库（基于 python-docx）。

封装「字体排版规范 v1 / 学术规范版」的全部排版细节：
- 中文走中文字体、英文数字走 Times New Roman，两者不混（中英混排不别扭的关键）。
- 三线表、页码、按文档大小决定的页眉。

典型用法见本文件底部 `if __name__ == "__main__"` 的最小示例，或技能 SKILL.md。
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 规范常量（学术规范版）------------------------------------------------
BODY_EAST = "宋体"           # 正文中文字体
HEAD_EAST = "黑体"           # 标题中文字体
WEST = "Times New Roman"     # 西文与数字字体
TITLE_PT = 18                # 大标题
H1_PT = 15                   # 一级标题
H2_PT = 13                   # 二级标题
BODY_PT = 12                 # 正文（小四）
TABLE_PT = 11                # 表内文字（比正文小一号）
PAGENO_PT = 10.5             # 页码（五号）
LINE = 1.5                   # 行距倍数

_MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'  # Word 公式(OMML)命名空间


def _set_run(run, size_pt, ascii_font=WEST, east_font=BODY_EAST, bold=False, color=None):
    """给一个文字片段设定中/西文字体、字号、加粗。"""
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), east_font)


def _set_line(p, multiple=LINE, before=0, after=0):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = multiple
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _set_outline(p, level0):
    """给标题段落写大纲级别（level0 从 0 起），让自动目录(TOC)能收录它。
    没有这一步，add_toc 生成的目录更新后会是空的。"""
    pPr = p._p.get_or_add_pPr()
    o = pPr.find(qn('w:outlineLvl'))
    if o is None:
        o = OxmlElement('w:outlineLvl'); pPr.append(o)
    o.set(qn('w:val'), str(level0))


def new_document():
    """新建 A4、2.5cm 页边距的空文档。"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Cm(2.5))
    return doc


def add_title(doc, zh, en=None):
    """大标题（黑体加粗居中）+ 可选英文副标题。"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(zh), TITLE_PT, east_font=HEAD_EAST, bold=True)
    _set_line(p)
    if en:
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(q.add_run(en), BODY_PT, east_font=HEAD_EAST)
        _set_line(q)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    """标题。level=1 一级（一、二…），level=2 二级（1.1…）。"""
    p = doc.add_paragraph()
    size = H1_PT if level == 1 else H2_PT
    _set_run(p.add_run(text), size, east_font=HEAD_EAST, bold=True)
    _set_line(p, before=6 if level == 1 else 4, after=4 if level == 1 else 2)
    _set_outline(p, 0 if level == 1 else 1)   # 关键：让标题进得了目录
    return p


def add_body(doc, runs):
    """正文段落，首行缩进 2 字符。

    runs: 字符串，或 [(文本, 是否加粗), ...] 列表（用于给关键词/数字加粗）。
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(BODY_PT * 2)
    if isinstance(runs, str):
        runs = [(runs, False)]
    for text, bold in runs:
        _set_run(p.add_run(text), BODY_PT, bold=bold)
    _set_line(p)
    return p


def add_bullets(doc, items):
    """无序列表（用 Word 的项目符号，不要手打 •）。"""
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        _set_run(p.add_run(it), BODY_PT)
        _set_line(p)


def _cell_border(cell, top=False, bottom=False):
    """只画上/下横线，竖线和中间线全去掉——三线表的关键。"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcb = tcPr.find(qn('w:tcBorders'))
    if tcb is None:
        tcb = OxmlElement('w:tcBorders'); tcPr.append(tcb)
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        e = tcb.find(qn('w:' + edge))
        if e is None:
            e = OxmlElement('w:' + edge); tcb.append(e)
        on = (edge == 'top' and top) or (edge == 'bottom' and bottom)
        if on:
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '8' if edge == 'top' else '6')
            e.set(qn('w:color'), '000000')
        else:
            e.set(qn('w:val'), 'nil')


def add_table_title(doc, text):
    p = doc.add_paragraph()
    _set_run(p.add_run(text), BODY_PT, east_font=HEAD_EAST, bold=True)
    _set_line(p, before=6)


def add_three_line_table(doc, data):
    """三线表。data 是二维列表，第一行为表头（自动加粗）。"""
    n = len(data)
    tbl = doc.add_table(rows=n, cols=len(data[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(data):
        for j, txt in enumerate(row):
            c = tbl.cell(i, j)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run(c.paragraphs[0].add_run(str(txt)), TABLE_PT, bold=(i == 0))
            _cell_border(c, top=(i == 0), bottom=(i == 0 or i == n - 1))
    return tbl


def add_figure(doc, image_path, caption=None, width_cm=14):
    """插入图片并在下方加图题（居中、黑体小一号）。图题在图下方是规范要求。"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Cm(width_cm))
    _set_line(p, before=6, after=2)
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(c.add_run(caption), TABLE_PT, east_font=HEAD_EAST, bold=True)
        _set_line(c, after=6)


def add_page_number(doc):
    """页脚居中页码，五号。固定加。"""
    sec = doc.sections[0]
    sec.footer.is_linked_to_previous = False
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_run(run, PAGENO_PT)
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE \\* MERGEFORMAT'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)


def add_header(doc, title):
    """页眉：居中文档标题 + 下方细横线。中等及以上文档才调用。"""
    sec = doc.sections[0]
    sec.header.is_linked_to_previous = False
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(title), PAGENO_PT)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '808080')
    pbdr.append(bottom); pPr.append(pbdr)


def add_toc(doc):
    """自动目录（列到二级标题）。打开文档后需按 F9 或确认更新域。长文档调用。"""
    p = doc.add_paragraph()
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    f3 = OxmlElement('w:t'); f3.text = "（打开文档后按 F9 更新目录）"
    f4 = OxmlElement('w:fldChar'); f4.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2, f3, f4):
        run._r.append(el)
    doc.add_paragraph()


def decide_layout(page_estimate, num_h1=0):
    """按文档大小返回是否加页眉/目录。

    返回 dict: {"header": bool, "toc": bool}
    短(≤5页): 都不加; 中(6-15页): 页眉加, 一级标题≥3 则目录加; 长(>15页): 都加。
    """
    if page_estimate <= 5:
        return {"header": False, "toc": False}
    if page_estimate <= 15:
        return {"header": True, "toc": num_h1 >= 3}
    return {"header": True, "toc": True}


_CN = {'零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4, '五': 5,
       '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}


def _cn2int(s):
    """中文数字转整数，支持 一..九十九。无法解析返回 None。"""
    if not s:
        return None
    if '十' in s:
        left, _, right = s.partition('十')
        tens = _CN.get(left, 1) if left else 1
        ones = _CN.get(right, 0) if right else 0
        return tens * 10 + ones
    if len(s) == 1 and s in _CN:
        return _CN[s]
    return None


def _h1_num(text):
    """取一级标题开头的序号（中文或阿拉伯），如『三、…』→3。取不到返回 None。"""
    t = text.strip()
    m = re.match(r'^第?\s*([一二三四五六七八九十两零]+)\s*[、.、)）章节]', t)
    if m:
        v = _cn2int(m.group(1))
        if v:
            return v
    m = re.match(r'^第?\s*(\d+)\s*[、.、)）章节]', t)
    if m:
        return int(m.group(1))
    return None


def _h2_num(text):
    """取二级标题开头的『a.b』编号，如『1.2 …』→(1, 2)。取不到返回 None。"""
    m = re.match(r'^\s*(\d+)\.(\d+)', text.strip())
    return (int(m.group(1)), int(m.group(2))) if m else None


def review(doc):
    """审查文档结构常见问题，返回警告字符串列表（空列表表示没发现问题）。

    检查项：
    - 重复子标题（同级同名标题出现多次）
    - 标题层级跳级（如从一级直接跳到三级）或未从一级开始
    - 标题序号不连续（一级『一、三、』缺号、二级『1.1 1.3』缺号、序号重复或倒退）
    - 二级标题的章节号与所属一级标题不符（如第二章下出现『3.1』）
    - 图/表编号不连续（缺号、重复、不从 1 开始）
    - 图或表的数量与图题/表题数量不匹配（有图无题、有题无图）
    - 交叉引用失效（正文写"见图5"却没有图5；图表定义了却从未被引用）
    - 公式编号不连续、公式交叉引用失效（"由式(5)"却无 (5) 号公式）
    - 公式未渲染（正文残留 \frac、$...$ 等 LaTeX 源码，未做成可编辑公式）
    - 空标题、把整句话当标题（以句末标点结尾）
    - 插了目录(TOC)却没有任何带层级的标题——更新后目录会为空
    建议在 doc.save() 之前调用：for w in ds.review(doc): print(w)
    """
    warnings = []
    headings = []  # [(级别从1起, 文本)]
    for p in doc.paragraphs:
        pPr = p._p.find(qn('w:pPr'))
        if pPr is None:
            continue
        o = pPr.find(qn('w:outlineLvl'))
        if o is None:
            continue
        headings.append((int(o.get(qn('w:val'))) + 1, p.text.strip()))

    # 重复子标题
    counts = {}
    for lvl, t in headings:
        if t:
            counts[(lvl, t)] = counts.get((lvl, t), 0) + 1
    for (lvl, t), c in counts.items():
        if c > 1:
            warnings.append(f"重复标题：{lvl} 级『{t}』出现 {c} 次")

    # 层级跳级 / 起始层级
    prev = 0
    for lvl, t in headings:
        if prev == 0 and lvl != 1:
            warnings.append(f"层级错误：第一个标题是 {lvl} 级，应从 1 级开始（『{t}』）")
        elif lvl - prev > 1:
            warnings.append(f"层级跳级：{prev} 级后直接出现 {lvl} 级（『{t}』），中间缺 {prev + 1} 级")
        prev = lvl

    # 标题序号连续性（标题里手写的序号）
    last_h1, last_h2, cur_chapter = 0, 0, None
    for lvl, t in headings:
        if lvl == 1:
            n = _h1_num(t)
            if n is not None:
                if n != last_h1 + 1:
                    if n <= last_h1:
                        warnings.append(f"序号问题：一级标题『{t}』序号 {n} 倒退或重复（上一个是 {last_h1}）")
                    else:
                        warnings.append(f"序号问题：一级标题序号从 {last_h1} 跳到 {n}，中间缺号（『{t}』）")
                last_h1, cur_chapter, last_h2 = n, n, 0
        elif lvl == 2:
            pair = _h2_num(t)
            if pair is not None:
                a, b = pair
                if cur_chapter is not None and a != cur_chapter:
                    warnings.append(f"序号问题：二级标题『{t}』的章节号 {a} 与所属第 {cur_chapter} 章不符")
                if b != last_h2 + 1:
                    if b <= last_h2:
                        warnings.append(f"序号问题：二级标题『{t}』序号 {a}.{b} 倒退或重复（上一个是 {a}.{last_h2}）")
                    else:
                        warnings.append(f"序号问题：二级标题序号 {a} 章内从 .{last_h2} 跳到 .{b}，中间缺号（『{t}』）")
                last_h2 = b

    # 标题质量：空标题、把整句话当标题
    for lvl, t in headings:
        if not t:
            warnings.append(f"空标题：存在一个 {lvl} 级标题没有文字")
        elif t[-1] in "。！？.!?":
            warnings.append(f"标题疑似整句：『{t}』以句末标点结尾，标题通常不带句号")

    # 图 / 表：编号连续性、与图题/表题数量是否匹配、与正文引用是否对应
    paras = [p.text.strip() for p in doc.paragraphs]
    cap_lines = {i for i, t in enumerate(paras) if re.match(r'^(图|表)\s*\d+', t)}

    def _caps(prefix):
        return [int(re.match(rf'^{prefix}\s*(\d+)', t).group(1))
                for t in paras if re.match(rf'^{prefix}\s*(\d+)', t)]

    def _refs(prefix):
        refs = set()
        for i, t in enumerate(paras):
            if i in cap_lines:        # 图题/表题本身不算引用
                continue
            for m in re.finditer(rf'{prefix}\s*(\d+)', t):
                refs.add(int(m.group(1)))
        return refs

    def _seq(nums, name):
        out, seen = [], set()
        for n in nums:
            if n in seen:
                out.append(f"{name}编号重复：{name}{n} 出现多次")
            seen.add(n)
        if seen:
            uniq = sorted(seen)
            if uniq[0] != 1:
                out.append(f"{name}编号未从 1 开始（第一个是 {name}{uniq[0]}）")
            missing = [x for x in range(1, uniq[-1] + 1) if x not in seen]
            if missing:
                out.append(f"{name}编号缺号：缺 " + "、".join(f"{name}{x}" for x in missing))
        return out

    for prefix, count in (("图", len(doc.inline_shapes)), ("表", len(doc.tables))):
        caps = _caps(prefix)
        warnings.extend(_seq(caps, prefix))
        if count > len(caps):
            warnings.append(f"{prefix}缺标题：文档有 {count} 个{prefix}，但只有 {len(caps)} 个{prefix}题")
        elif len(caps) > count:
            warnings.append(f"{prefix}题多余：有 {len(caps)} 个{prefix}题，却只有 {count} 个{prefix}")
        defined, refs = set(caps), _refs(prefix)
        for r in sorted(refs - defined):
            warnings.append(f"引用失效：正文提到『{prefix}{r}』，但没有对应的{prefix}题")
        if refs:   # 文档用了"见图N"式引用，才提示哪些没被引用
            for d in sorted(defined - refs):
                warnings.append(f"提示：{prefix}{d} 在正文中没有被引用")

    # 公式：编号连续性、交叉引用、是否为可编辑公式（未渲染的 LaTeX 残留）
    eq_def, label_idx = [], set()
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        has_math = p._p.find('.//{%s}oMath' % _MATH_NS) is not None
        m = re.search(r'[\(（](\d+)[\)）]\s*$', t)   # 行尾的 (N)，公式编号惯例
        if m and (has_math or p.alignment == WD_ALIGN_PARAGRAPH.RIGHT):
            eq_def.append(int(m.group(1)))
            label_idx.add(i)
    eq_ref = set()
    for i, t in enumerate(paras):
        if i in label_idx:
            continue
        for mm in re.finditer(r'(?:式|公式|方程)\s*[\(（](\d+)[\)）]', t):
            eq_ref.add(int(mm.group(1)))
    warnings.extend(_seq(eq_def, '式'))
    for r in sorted(eq_ref - set(eq_def)):
        warnings.append(f"引用失效：正文提到『式({r})』，但找不到编号为 ({r}) 的公式")
    if eq_ref and eq_def:
        for d in sorted(set(eq_def) - eq_ref):
            warnings.append(f"提示：公式({d}) 没有在正文中被引用")

    _tex = re.compile(
        r'\\(?:frac|dfrac|sum|prod|int|sqrt|left|right|begin|end|cdot|times|div|'
        r'leq|geq|neq|approx|partial|nabla|infty|alpha|beta|gamma|delta|theta|'
        r'lambda|sigma|omega|mu|pi|mathrm|mathbf|mathbb|text)\b|\^\{|_\{|\$[^$]+\$')
    tex_hits = sorted({t for t in paras if _tex.search(t)})
    if tex_hits:
        warnings.append(
            f"公式可能未渲染：发现 {len(tex_hits)} 处疑似 LaTeX 源码"
            f"（如『{tex_hits[0][:30]}…』），应作为可编辑公式插入，而非留作纯文本")

    # 目录是否会空
    has_toc = any(
        (r.text or "").strip().startswith("TOC")
        for p in doc.paragraphs
        for r in p._p.findall('.//' + qn('w:instrText'))
    )
    if has_toc and not headings:
        warnings.append("目录错误：插入了目录(TOC)，但没有任何带层级的标题，更新后目录将为空")

    return warnings


# ===========================================================================
#  字体强制修复 enforce_fonts —— 遍历每个 run 重设中西文字体
#  解决「改已有文档 / 手写 run 时中文回退到默认字体」这个高频坑。
# ===========================================================================

_CJK = re.compile(r'[㐀-鿿豈-﫿　-〿＀-￯]')


def _has_cjk(s):
    """字符串里是否含中文（含中文标点、全角符号）。"""
    return bool(_CJK.search(s or ""))


def _heading_level(p):
    """段落是否标题、第几级（1 起）。靠大纲级别判断，取不到再看样式名。无则 None。"""
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        o = pPr.find(qn('w:outlineLvl'))
        if o is not None:
            try:
                return int(o.get(qn('w:val'))) + 1
            except (TypeError, ValueError):
                pass
    name = (p.style.name if p.style is not None else "") or ""
    if name.startswith("Heading"):
        tail = name.split()[-1]
        return int(tail) if tail.isdigit() else 1
    return None


def _run_east(run):
    """读 run 已设的 eastAsia（中文）字体名，没设返回 None。"""
    rpr = run._element.find(qn('w:rPr'))
    if rpr is None:
        return None
    rf = rpr.find(qn('w:rFonts'))
    return rf.get(qn('w:eastAsia')) if rf is not None else None


def _force_run_fonts(run, east):
    """只改字体（西文 Times、中文 east），不动字号/加粗。"""
    run.font.name = WEST
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), WEST)
    rf.set(qn('w:hAnsi'), WEST)
    rf.set(qn('w:eastAsia'), east)


def _block_paragraphs(doc, with_hf=True):
    """文档正文 + 所有表格单元格 +（可选）页眉页脚里的全部段落，标注来源。

    产出 (kind, paragraph)，kind ∈ {'body','table','header','footer'}。
    """
    for p in doc.paragraphs:
        yield 'body', p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield 'table', p
    if with_hf:
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                yield 'header', p
            for p in sec.footer.paragraphs:
                yield 'footer', p


def enforce_fonts(doc, body_east=BODY_EAST, head_east=HEAD_EAST):
    """遍历文档每一个 run，强制「西文走 Times、中文走宋体/黑体」。

    用于：改一份已有 docx、或别处手写了 run 之后兜底修字体——这些都绕过了 add_* 函数，
    中文容易回退到默认字体。本函数不分谁加的，把每个 run 的字体重设到位。
    标题段（有大纲级别或 Heading 样式）用黑体，其余用宋体。**只改字体，不动字号和加粗。**

    返回被改动的 run 数。
    """
    n = 0
    for _kind, p in _block_paragraphs(doc):
        east = head_east if _heading_level(p) is not None else body_east
        for run in p.runs:
            _force_run_fonts(run, east)
            n += 1
    return n


# ===========================================================================
#  合规体检 lint —— 对任意 .docx 逐条打分（规范 spec.md 的 11 条硬规则）
#  与 review() 互补：review 查「结构」（标题层级/序号/图表引用），lint 查「格式」。
# ===========================================================================

def _approx(a, b, tol):
    return a is not None and abs(a - b) <= tol


def _is_toc_para(p):
    """是否目录(TOC)域所在段落——这种段落不算正文。"""
    for it in p._p.findall('.//' + qn('w:instrText')):
        if (it.text or "").strip().startswith("TOC"):
            return True
    return False


def _is_list_para(p):
    """是否项目符号/编号列表段——它们是悬挂缩进，不该按首行缩进考核。"""
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:numPr')) is not None:
        return True
    name = (p.style.name if p.style is not None else "") or ""
    return name.startswith("List")


def _mode(vals):
    """众数（出现最多的值）；空则 None。"""
    vals = [v for v in vals if v is not None]
    if not vals:
        return None
    best, bestc = None, -1
    for v in set(vals):
        c = vals.count(v)
        if c > bestc:
            best, bestc = v, c
    return best


def _footer_has_page(doc):
    for sec in doc.sections:
        for it in sec.footer._element.iter(qn('w:instrText')):
            if 'PAGE' in (it.text or ""):
                return True
    return False


def _has_toc_field(doc):
    for p in doc.paragraphs:
        if _is_toc_para(p):
            return True
    return False


def _header_has_text(doc):
    for sec in doc.sections:
        if any((p.text or "").strip() for p in sec.header.paragraphs):
            return True
    return False


def _table_is_three_line(tbl):
    """三线表判定：用了 Table Grid 样式、或任一单元格有竖线(左/右/中竖)，即判否。"""
    sty = tbl.style.name if tbl.style is not None else ""
    if sty and sty.replace(" ", "").lower() == "tablegrid":
        return False
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                continue
            tcb = tcPr.find(qn('w:tcBorders'))
            if tcb is None:
                continue
            for edge in ('left', 'right', 'insideV'):
                e = tcb.find(qn('w:' + edge))
                if e is not None and (e.get(qn('w:val')) not in (None, 'nil', 'none')):
                    return False
    return True


def lint(doc, english=None):
    """对已打开的 doc 做格式合规体检（规范的 11 条硬规则）。

    english: True=按纯英文文档（跳过中文字体相关项）；None=按内容自动判断。
    返回 dict：
      {"english": bool, "score": 通过数, "total": 计分项数,
       "skipped": 跳过数, "items": [ {id, name, status, expected, actual}, ... ]}
    status ∈ {"PASS","FAIL","SKIP"}。SKIP 表示该项不适用（如纯英文跳过中文字体、
    短文档不要求目录），不计入分母。
    """
    items = []

    def add(rid, name, status, expected="", actual=""):
        items.append({"id": rid, "name": name, "status": status,
                      "expected": str(expected), "actual": str(actual)})

    # ---- 预收集 ----
    blocks = list(_block_paragraphs(doc))
    # 正文段：非标题、非空、非目录、非居中（排除大标题/图表题）
    body = [p for k, p in blocks if k == 'body' and p.text.strip()
            and _heading_level(p) is None and not _is_toc_para(p)
            and p.alignment != WD_ALIGN_PARAGRAPH.CENTER]
    headings = [(_heading_level(p), p) for k, p in blocks if k == 'body'
                and _heading_level(p) is not None]
    h1s = [p for lvl, p in headings if lvl == 1]

    # 自动判断中英：统计所有 run 文本里中文字符占比
    if english is None:
        cjk = sum(len(_CJK.findall(r.text or "")) for _k, p in blocks for r in p.runs)
        letters = sum(sum(c.isascii() and c.isalpha() for c in (r.text or ""))
                      for _k, p in blocks for r in p.runs)
        english = (cjk == 0) or (letters > 0 and cjk / (cjk + letters) < 0.05)

    sec = doc.sections[0]

    # ---- 规则 1：正文字号 12pt ----
    sizes = [r.font.size.pt for p in body for r in p.runs if r.font.size is not None]
    bs = _mode(sizes)
    if bs is None:
        normal = doc.styles['Normal'].font.size
        bs = normal.pt if normal is not None else None
    add(1, "正文字号 12pt", "PASS" if _approx(bs, 12, 0.3) else "FAIL",
        "12pt", f"{bs}pt" if bs is not None else "未设(继承)")

    # ---- 规则 2：行距 1.5 倍 ----
    ls = _mode([p.paragraph_format.line_spacing for p in body
                if p.paragraph_format.line_spacing is not None])
    add(2, "行距 1.5 倍", "PASS" if _approx(ls, 1.5, 0.01) else "FAIL",
        "1.5", ls if ls is not None else "未设(默认单倍)")

    # ---- 规则 3：标题中文字体 黑体 ----
    if english:
        add(3, "标题中文字体 黑体", "SKIP", "黑体", "纯英文，跳过")
    elif not headings:
        add(3, "标题中文字体 黑体", "SKIP", "黑体", "无带大纲级别的标题")
    else:
        he = [_run_east(r) for lvl, p in headings for r in p.runs if _has_cjk(r.text)]
        he = [x for x in he if x]
        ok = he and all(x == HEAD_EAST for x in he)
        add(3, "标题中文字体 黑体", "PASS" if ok else "FAIL",
            "黑体", _mode(he) or "未设")

    # ---- 规则 4：一级标题字号 15pt ----
    if not h1s:
        add(4, "一级标题字号 15pt", "SKIP", "15pt", "无一级标题")
    else:
        h1sz = _mode([r.font.size.pt for p in h1s for r in p.runs if r.font.size is not None])
        add(4, "一级标题字号 15pt", "PASS" if _approx(h1sz, 15, 0.3) else "FAIL",
            "15pt", f"{h1sz}pt" if h1sz is not None else "未设")

    # ---- 规则 5：表格三线表 ----
    if not doc.tables:
        add(5, "表格三线表", "SKIP", "三线表", "无表格")
    else:
        bad = [i + 1 for i, t in enumerate(doc.tables) if not _table_is_three_line(t)]
        add(5, "表格三线表", "PASS" if not bad else "FAIL",
            "仅顶/底/表头下三线", f"全部 {len(doc.tables)} 张合规" if not bad
            else f"{len(bad)}/{len(doc.tables)} 张含竖线或用了 Table Grid")

    # ---- 规则 6：页码 ----
    add(6, "页脚页码", "PASS" if _footer_has_page(doc) else "FAIL",
        "页脚有 PAGE 域", "有" if _footer_has_page(doc) else "无页码")

    # ---- 规则 7：页面 A4 + 2.5cm ----
    pw, ph = sec.page_width.cm, sec.page_height.cm
    margins = [sec.top_margin.cm, sec.bottom_margin.cm, sec.left_margin.cm, sec.right_margin.cm]
    page_ok = _approx(pw, 21.0, 0.2) and _approx(ph, 29.7, 0.3) \
        and all(_approx(m, 2.5, 0.2) for m in margins)
    add(7, "页面 A4 + 2.5cm 边距", "PASS" if page_ok else "FAIL",
        "21×29.7cm, 边距2.5cm", f"{pw:.1f}×{ph:.1f}cm, 边距{'/'.join(f'{m:.1f}' for m in margins)}cm")

    # ---- 规则 8：正文首行缩进 2 字符 ----
    subst = [p for p in body if len(p.text.strip()) >= 15 and not _is_list_para(p)
             and p.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY)]
    if not subst:
        add(8, "正文首行缩进 2 字符", "SKIP", "≈24pt", "无成段正文")
    else:
        good = sum(1 for p in subst
                   if p.paragraph_format.first_line_indent is not None
                   and _approx(p.paragraph_format.first_line_indent.pt, 24, 4))
        add(8, "正文首行缩进 2 字符", "PASS" if good >= len(subst) * 0.5 else "FAIL",
            "≈24pt", f"{good}/{len(subst)} 段有缩进")

    # ---- 规则 9：一级标题段前 6 / 段后 4 ----
    if not h1s:
        add(9, "一级标题段前后距", "SKIP", "段前6/段后4", "无一级标题")
    else:
        ok = all((p.paragraph_format.space_before is not None and p.paragraph_format.space_before.pt > 0)
                 and (p.paragraph_format.space_after is not None and p.paragraph_format.space_after.pt > 0)
                 for p in h1s)
        add(9, "一级标题段前后距", "PASS" if ok else "FAIL",
            "段前6pt/段后4pt", "已设" if ok else "缺段前或段后")

    # ---- 规则 10：长文档目录 + 页眉 ----
    pages = max(1, round((len(body) + 2 * len(doc.tables)) / 12))
    layout = decide_layout(pages, len(h1s))
    need = []
    if layout["header"]:
        need.append(("页眉", _header_has_text(doc)))
    if layout["toc"]:
        need.append(("目录", _has_toc_field(doc)))
    if not need:
        add(10, "篇幅对应的页眉/目录", "SKIP", "短文档不要求", f"约 {pages} 页")
    else:
        miss = [nm for nm, ok in need if not ok]
        add(10, "篇幅对应的页眉/目录", "PASS" if not miss else "FAIL",
            f"约{pages}页应有 " + "+".join(nm for nm, _ in need),
            "齐全" if not miss else "缺 " + "+".join(miss))

    # ---- 规则 11：中文 run 都带 eastAsia（不回退）----
    if english:
        add(11, "中文 run 均设字体(不回退)", "SKIP", "每个中文片段带 eastAsia", "纯英文，跳过")
    else:
        miss = 0
        for _k, p in blocks:
            for r in p.runs:
                if _has_cjk(r.text) and not _run_east(r):
                    miss += 1
        add(11, "中文 run 均设字体(不回退)", "PASS" if miss == 0 else "FAIL",
            "每个中文片段带 eastAsia", "全部已设" if miss == 0 else f"{miss} 个中文片段未设字体")

    score = sum(1 for it in items if it["status"] == "PASS")
    total = sum(1 for it in items if it["status"] != "SKIP")
    skipped = sum(1 for it in items if it["status"] == "SKIP")
    return {"english": english, "score": score, "total": total,
            "skipped": skipped, "items": items}


if __name__ == "__main__":
    # 最小示例 + 自检：用本库生成的文档，lint 应当满分
    doc = new_document()
    add_page_number(doc)
    add_title(doc, "示例文档", "A Sample Document")
    add_heading(doc, "一、引言", 1)
    add_body(doc, [("正文里数字 ", False), ("2.9%", True), (" 会加粗。", False)])
    add_three_line_table(doc, [["列A", "列B"], ["1", "2"], ["3", "4"]])
    doc.save("示例.docx")
    rep = lint(doc)
    print("已生成 示例.docx；自检得分 %d/%d（跳过 %d）" % (rep["score"], rep["total"], rep["skipped"]))
    for it in rep["items"]:
        print(" ", it["status"], it["name"], "—", it["actual"])
